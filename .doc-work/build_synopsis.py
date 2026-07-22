from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path(r"C:\Users\pc\Downloads\Travelogue Studio\Synopsys Format - Working Copy.docx")
OUTPUT = Path(r"C:\Users\pc\Downloads\Travelogue Studio\Yaara Photobook Project Synopsis.docx")
TITLE = "YAARA - TRAVELOGUE STUDIO: A WEB-BASED PHOTOBOOK EDITOR AND PREVIEW PLATFORM"


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def replace_text(paragraph, text, size=12, bold=False):
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def insert_after(paragraph, new_paragraph):
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def new_paragraph_like(doc, anchor, text="", style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if text:
        p.add_run(text)
    insert_after(anchor, p)
    return p


def format_body(p):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.35)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.widow_control = True
    for run in p.runs:
        set_run_font(run, 12)


def format_heading(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(run, 14, bold=True)


def add_body(doc, anchor, text):
    p = new_paragraph_like(doc, anchor, text)
    format_body(p)
    return p


def ensure_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall("w:abstractNum", ns)]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall("w:num", ns)]
    abstract_id = max(abs_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), "bullet"); lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), "•"); lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc"); lvl_jc.set(qn("w:val"), "left"); lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "504"); tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "504"); ind.set(qn("w:hanging"), "288"); ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr"); fonts = OxmlElement("w:rFonts"); fonts.set(qn("w:ascii"), "Arial"); fonts.set(qn("w:hAnsi"), "Arial"); rpr.append(fonts); lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId"); abstract_ref.set(qn("w:val"), str(abstract_id)); num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, anchor, text, num_id):
    p = new_paragraph_like(doc, anchor, text)
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); num_pr.append(ilvl)
    num = OxmlElement("w:numId"); num.set(qn("w:val"), str(num_id)); num_pr.append(num)
    ppr.append(num_pr)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    for run in p.runs:
        set_run_font(run, 11.5)
    return p


doc = Document(SOURCE)

# Preserve the supplied cover while filling the two project-title slots.
replace_text(doc.paragraphs[3], TITLE, size=18, bold=True)
doc.paragraphs[3].paragraph_format.line_spacing = 1.05
replace_text(doc.paragraphs[15], TITLE.title(), size=16, bold=True)
doc.paragraphs[15].paragraph_format.line_spacing = 1.05

# Normalize the supplied submission table without inventing personal details.
table = doc.tables[0]
table.autofit = False
for row in table.rows:
    for cell in row.cells:
        cell.width = Inches(3.2)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                set_run_font(r, 11, bold=("Submitted" in r.text))

# Keep the supplied content list, then begin the populated synopsis on a new page.
for idx in range(17, 23):
    p = doc.paragraphs[idx]
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        set_run_font(r, 12, bold=(idx == 17))

anchor = doc.paragraphs[22]
page_break = new_paragraph_like(doc, anchor)
page_break.add_run().add_break(WD_BREAK.PAGE)
anchor = page_break
bullet_num_id = ensure_bullet_numbering(doc)

sections = [
    ("1. Problem Statement", [
        "Photographs are captured in large numbers, but converting them into a coherent, printable travel photobook remains a time-consuming task. Users often have to switch between file managers, generic design tools, and PDF utilities to select images, arrange layouts, add captions, preview pages, and prepare an export. These tools usually demand design experience and make it difficult to maintain visual consistency across an entire book.",
        "The project addresses this gap through Yaara (Travelogue Studio), a browser-based photobook creation system focused on a short, practical workflow: choose templates, upload photographs, fill pages, customize the design, preview the book, and export it. The system must protect administrator-authored templates while still allowing users to replace photographs and edit permitted content. It must also preserve user intent during automation: already-filled frames are not overwritten, excluded photographs remain excluded, and empty locked frames are unlocked only when a fill action needs them.",
        "A further challenge is reliable project continuity. Photographs and project data need local persistence during editing, while production template metadata and shared media require secure cloud storage. The same editable project title and stable template identifiers must remain consistent in the editor, preview, PDF export, and .wanderbook project-file flows."
    ]),
    ("2. Objectives", [
        "Provide a fast, template-first workflow for creating a complete square photobook from uploaded travel photographs.",
        "Allow users to select one or more categorized templates, preserve their chosen order, and open the resulting pages directly in the editor.",
        "Support photo upload, drag-and-drop placement, crop, zoom, rotation, filters, frames, shape masks, captions, stickers, text, drawings, and customizable page backgrounds.",
        "Implement Magic Fill to populate empty frames across the whole book without replacing existing photographs or using excluded images.",
        "Protect administrator-created page structures, backgrounds, and locked elements while keeping user-created layouts flexible.",
        "Provide undo/redo, page management, responsive editing controls, realistic book preview, shareable preview support, PDF export, and compatible project saving/loading.",
        "Use secure production persistence in which Supabase stores structured metadata and ImageKit stores media, with private credentials restricted to server-side functions."
    ]),
    ("3. Proposed Methodology", [
        "The system follows a component-based, state-driven web architecture. TanStack Start supplies routing and server-function support, while React and TypeScript implement the landing page, editor, preview, administration, and reusable photobook components. The application state is centralized in a Zustand store and augmented with Zundo history so editing commands can participate in undo and redo.",
        "The user begins on a compact template-discovery interface. Templates are loaded by category, previewed in a square card, and added to an ordered selection bucket. When the user proceeds, the selected template data is converted into book pages while retaining stable template IDs and protection metadata. Photographs are then uploaded into a library, with binary image data persisted locally through IndexedDB and lightweight editor state persisted through browser storage.",
        "During editing, each page is represented as structured data containing a background and typed elements such as photos, stickers, text, quotations, and drawings. Photo elements retain position, size, rotation, crop offsets, zoom, opacity, filter settings, frame style, shape mask, captions, lock state, and optional background-removal or magic-mask data. The editor renders these elements on a square 5.5-inch page model and exposes page, design, library, and element-level controls.",
        "Automation is implemented as deterministic store actions. Magic Fill scans the entire book for empty photo frames, selects only non-excluded library images, keeps existing frame assignments unchanged, and unlocks a locked frame only when that empty frame is filled. Magic Layout can derive a masked photo slot from a template or background image. Client-side subject segmentation supports background removal, with erase and restore refinement for the resulting mask.",
        "For output, the preview route renders the pages as an interactive page-flipping book. PDF export uses jsPDF and canvas rendering to reproduce page backgrounds, crop and transform settings, masks, filters, frames, text, stickers, drawings, captions, and overlays. Production administration and shared-preview functions use server-side APIs; Supabase holds structured records and ImageKit holds uploaded media so private keys never reach the browser.",
        "Verification is carried out through TypeScript compilation, linting, production builds, and manual responsive checks at desktop and narrow mobile widths. Functional tests focus on template selection order, protected-template behavior, upload persistence, Magic Fill safeguards, undo/redo, project compatibility, preview accuracy, and PDF output."
    ]),
    ("4. Technologies/Tools to be Used", [
        "Frontend framework: React 19 with TanStack Start, TanStack Router, and Vite.",
        "Programming language and validation: TypeScript and Zod.",
        "User interface: Tailwind CSS 4, Radix UI primitives, lucide-react icons, and responsive CSS.",
        "State and history: Zustand for centralized photobook state and Zundo for undo/redo history.",
        "Interactive editing: dnd-kit for drag-and-drop, react-rnd and resizable panels for positioning and resizing, and browser Canvas APIs for rendering and masks.",
        "Preview and export: react-pageflip for book presentation, html2canvas-pro where DOM capture is required, and jsPDF for downloadable PDF generation.",
        "Image processing: @imgly/background-removal for client-side foreground segmentation, with custom crop, filter, shape-mask, and compositing logic.",
        "Persistence and deployment: IndexedDB and browser storage for local work; Supabase for production metadata; ImageKit for production media; Vercel for deployment.",
        "Development quality tools: ESLint, Prettier, TypeScript compiler, npm, and Git."
    ]),
    ("5. Expected Outcomes", [
        "A working responsive web application that enables users to create a polished travel photobook from template selection through final export.",
        "A streamlined editor that supports multi-page books, rich photo styling, text and decorative elements, backgrounds, drawings, project titles, and reusable templates.",
        "Reliable automated filling that saves time while respecting filled frames, excluded photographs, and template locks.",
        "A realistic interactive preview and a high-quality PDF whose page content closely matches the editor, including crop, masks, filters, frames, stickers, text, and drawings.",
        "Persistent projects that can be resumed locally and remain compatible with saved .wanderbook files and stable template identifiers.",
        "A protected administration workflow for publishing and maintaining templates, categories, backgrounds, stickers, overlays, and thumbnails without exposing private cloud credentials.",
        "An extensible architecture that can later support collaboration, print-service integration, additional book sizes, advanced asset search, and richer sharing controls."
    ]),
]

for section_index, (heading, items) in enumerate(sections):
    hp = new_paragraph_like(doc, anchor, heading)
    if section_index > 0:
        hp.paragraph_format.page_break_before = True
    format_heading(hp)
    anchor = hp
    if section_index in (1, 3, 4):
        for item in items:
            anchor = add_bullet(doc, anchor, item, bullet_num_id)
    else:
        for item in items:
            anchor = add_body(doc, anchor, item)

# Add simple footer page numbers while retaining all existing section settings.
section = doc.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Page ")
set_run_font(run, 10)
fld_char1 = OxmlElement("w:fldChar")
fld_char1.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText")
instr.set(qn("xml:space"), "preserve")
instr.text = " PAGE "
fld_char2 = OxmlElement("w:fldChar")
fld_char2.set(qn("w:fldCharType"), "end")
run._r.append(fld_char1)
run._r.append(instr)
run._r.append(fld_char2)

doc.core_properties.title = "Yaara Photobook Project Synopsis"
doc.core_properties.subject = "Bachelor of Technology project synopsis"
doc.core_properties.keywords = "Yaara, Travelogue Studio, photobook editor, React, TanStack Start"
doc.save(OUTPUT)
print(OUTPUT)
